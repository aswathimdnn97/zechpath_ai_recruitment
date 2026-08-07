from docx import Document
def read_docx(file):
    """Read a DOCX resume and return its text."""
    document=Document(file)
    
    text=""
    #read paragraph
    for paragraph in document.paragraphs:
        if(paragraph.text.strip()!=""):
            text=text+paragraph.text +"\n"

    #read table
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if(cell.text.strip!=""):
                    text=text+cell.text+" "
        text=text+"\n"
        
    return text